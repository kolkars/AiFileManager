from pathlib import Path

from docx import Document

from .base import ContentUnit, ExtractionResult, normalize_text


class DocxExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        document = Document(path)
        units: list[ContentUnit] = []
        parts: list[str] = []
        for index, paragraph in enumerate(document.paragraphs, 1):
            value = normalize_text(paragraph.text)
            if not value:
                continue
            kind = "heading" if paragraph.style and paragraph.style.name.startswith("Heading") else "paragraph"
            units.append(ContentUnit(kind, len(units), value, f"paragraph:{index}", {"paragraph": index, "style": paragraph.style.name if paragraph.style else None}))
            parts.append(value)
        for table_number, table in enumerate(document.tables, 1):
            value = normalize_text("\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows))
            units.append(ContentUnit("table", len(units), value, f"table:{table_number}", {"table": table_number, "rows": len(table.rows), "columns": len(table.columns)}))
            parts.append(value)
        return ExtractionResult(normalize_text("\n\n".join(parts)), tuple(units), {"format": "docx", "paragraph_count": len(document.paragraphs), "table_count": len(document.tables)})
